import os
import openai
import tkinter as tk
from tkinter import scrolledtext, simpledialog, messagebox
from docx import Document

# Obter a chave da API da variável de ambiente
api_key = os.getenv('API_KEY')
openai.api_key = api_key

class ChatGPTApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ChatGPT Clone")
        self.root.geometry("600x600")

        self.messages = [
            {"role": "system", "content": "You are a helpful assistant."}
        ]

        self.create_widgets()

    def create_widgets(self):
        self.chat_window = scrolledtext.ScrolledText(self.root, wrap=tk.WORD)
        self.chat_window.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        self.chat_window.config(state=tk.DISABLED)

        self.input_frame = tk.Frame(self.root)
        self.input_frame.pack(pady=10, padx=10, fill=tk.X, expand=False)

        self.entry = tk.Entry(self.input_frame, width=80)
        self.entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        self.entry.bind("<Return>", self.send_message)

        self.send_button = tk.Button(self.input_frame, text="Send", command=self.send_message)
        self.send_button.pack(side=tk.RIGHT)

        self.doc_button = tk.Button(self.root, text="Generate Academic Paper", command=self.generate_academic_paper)
        self.doc_button.pack(pady=10)

    def send_message(self, event=None):
        user_message = self.entry.get()
        if user_message.strip():
            self.append_message("user", user_message)
            self.entry.delete(0, tk.END)
            self.get_response(user_message)

    def append_message(self, role, content):
        self.messages.append({"role": role, "content": content})
        self.chat_window.config(state=tk.NORMAL)
        self.chat_window.insert(tk.END, f"{role.capitalize()}: {content}\n")
        self.chat_window.config(state=tk.DISABLED)
        self.chat_window.yview(tk.END)

    def get_response(self, user_message):
        try:
            client = openai.OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=self.messages
            )
            assistant_message = response.choices[0].message.content
            self.append_message("assistant", assistant_message)
        except Exception as e:
            self.append_message("system", f"Error: {str(e)}")

    def generate_academic_paper(self):
        theme = simpledialog.askstring("Input", "Enter the theme for the academic paper:")
        if not theme:
            messagebox.showerror("Error", "Theme is required!")
            return
        
        prompt = f"Write an academic paper on the theme '{theme}'. Divide the text into categories: Title, Subtitle, and Paragraphs."
        self.append_message("user", prompt)
        self.get_response(prompt)
        
        try:
            client = openai.OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=self.messages
            )
            paper_content = response.choices[0].message.content
            self.create_docx(paper_content)
            messagebox.showinfo("Success", "Academic paper generated successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate paper: {str(e)}")

    def create_docx(self, content):
        document = Document()
        lines = content.split("\n")
        for line in lines:
            if line.startswith("Title:"):
                document.add_heading(line.replace("Title:", "").strip(), level=1)
            elif line.startswith("Subtitle:"):
                document.add_heading(line.replace("Subtitle:", "").strip(), level=2)
            else:
                document.add_paragraph(line.strip())

        document.save("academic_paper.docx")

if __name__ == "__main__":
    root = tk.Tk()
    app = ChatGPTApp(root)
    root.mainloop()
