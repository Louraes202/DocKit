import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font_style(font, size=None, bold=None, italic=None, name=None, color=None):
    if size:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if name:
        font.name = name
    if color:
        font.color.rgb = RGBColor(color[0], color[1], color[2])

def apply_custom_styles(styles, custom_styles):
    for style_name, style_attributes in custom_styles.items():
        if style_name in styles:
            style = styles[style_name]
            set_font_style(
                style.font,
                size=style_attributes.get('size'),
                bold=style_attributes.get('bold'),
                italic=style_attributes.get('italic'),
                name=style_attributes.get('name'),
                color=style_attributes.get('color')
            )

def add_paragraph(doc, text, style=None, alignment=None):
    paragraph = doc.add_paragraph(text, style=style)
    if alignment:
        if alignment == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return paragraph

def add_image(doc, image):
    run = doc.add_paragraph().add_run()
    run.add_picture(image['path'], width=Inches(image['width']), height=Inches(image['height']))
    if 'caption' in image:
        caption_paragraph = doc.add_paragraph(image['caption'], style='Caption')
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_table(doc, table_data):
    table = doc.add_table(rows=1, cols=len(table_data['headers']))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(table_data['headers']):
        hdr_cells[idx].text = header

    for row in table_data['rows']:
        row_cells = table.add_row().cells
        for idx, cell in enumerate(row):
            row_cells[idx].text = cell

def add_list(doc, items, list_type='bullet'):
    if list_type == 'bullet':
        for item in items:
            doc.add_paragraph(item, style='ListBullet')
    elif list_type == 'number':
        for item in items:
            doc.add_paragraph(item, style='ListNumber')

def add_content(doc, content):
    doc.add_heading(content['title'], level=0)

    def add_section(section, level=1):
        doc.add_heading(section['subtitle'], level=level)
        for paragraph in section.get('paragraphs', []):
            add_paragraph(doc, paragraph['text'], style='BodyText', alignment=paragraph.get('alignment'))
        for image in section.get('images', []):
            add_image(doc, image)
        for table in section.get('tables', []):
            add_table(doc, table)
        for list_data in section.get('lists', []):
            add_list(doc, list_data['items'], list_type=list_data.get('type', 'bullet'))
        for sub_section in section.get('subsections', []):
            add_section(sub_section, level=level+1)

    for section in content['sections']:
        add_section(section)

def generate_document(content, theme=None, custom_styles=None):
    doc = Document()

    if theme:
        apply_theme(doc.styles, theme)

    if custom_styles:
        apply_custom_styles(doc.styles, custom_styles)

    add_content(doc, content)

    document_path = 'output.docx'
    doc.save(document_path)
    return document_path

if __name__ == "__main__":
    with open('content.json', 'r') as f:
        data = json.load(f)
        generate_document(data['content'], data.get('theme'), data.get('custom_styles'))