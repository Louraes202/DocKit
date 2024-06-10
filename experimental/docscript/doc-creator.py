from docx import Document
from docx.shared import Pt

def create_document(content, themes):
    # Cria um novo documento
    doc = Document()
    
    # Define o estilo do documento
    styles = doc.styles
    title_style = styles['Title']
    subtitle_style = styles['Heading 1']
    body_style = styles['Normal']

    # Personaliza estilos (opcional)
    title_font = title_style.font
    title_font.size = Pt(24)

    subtitle_font = subtitle_style.font
    subtitle_font.size = Pt(18)

    body_font = body_style.font
    body_font.size = Pt(12)

    # Adiciona título
    doc.add_heading(content['title'], level=0)

    # Adiciona subtítulos e parágrafos
    for section in content['sections']:
        doc.add_heading(section['subtitle'], level=1)
        for paragraph in section['paragraphs']:
            doc.add_paragraph(paragraph, style='BodyText')

    # Salva o documento
    doc.save('output.docx')

# Exemplo de conteúdo
content = {
    'title': 'Exemplo de Documento',
    'sections': [
        {
            'subtitle': 'Introdução',
            'paragraphs': [
                'Este é um parágrafo introdutório.',
                'Este é outro parágrafo na introdução.'
            ]
        },
        {
            'subtitle': 'Desenvolvimento',
            'paragraphs': [
                'Este é um parágrafo no desenvolvimento.',
                'Este é outro parágrafo no desenvolvimento.'
            ]
        },
        {
            'subtitle': 'Conclusão',
            'paragraphs': [
                'Este é um parágrafo na conclusão.',
                'Este é outro parágrafo na conclusão.'
            ]
        }
    ]
}

# Chama a função para criar o documento
create_document(content, None)
