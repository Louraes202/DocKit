import openai
import os
from docx import Document
from docx.shared import Pt

# Obter a chave da API da variável de ambiente
api_key = os.getenv('API_KEY')
openai.api_key = api_key

def generate_content_from_prompt(prompt, max_tokens):
    client = openai.OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Retorna uma prompt informativa de acordo com o que o utilizador te está a pedir e com o formato pedido. Cada secção deve ter a informação necessária da pesquisa. Não retornes resposta nenhuma a esta mensagem."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=max_tokens
    )
    return response.choices[0].message.content.strip()

def create_document(content, theme):
    doc = Document()

    styles = doc.styles
    title_style = styles['Title']
    subtitle_style = styles['Heading 1']
    body_style = styles['Normal']

    if theme == 'corporate':
        title_font = title_style.font
        title_font.size = Pt(26)
        title_font.bold = True

        subtitle_font = subtitle_style.font
        subtitle_font.size = Pt(20)
        subtitle_font.bold = True

        body_font = body_style.font
        body_font.size = Pt(12)
        body_font.name = 'Arial'

    elif theme == 'academic':
        title_font = title_style.font
        title_font.size = Pt(24)
        title_font.italic = True

        subtitle_font = subtitle_style.font
        subtitle_font.size = Pt(18)
        subtitle_font.italic = True

        body_font = body_style.font
        body_font.size = Pt(11)
        body_font.name = 'Times New Roman'

    # Adiciona título
    doc.add_heading(content['title'], level=0)

    # Adiciona subtítulos e parágrafos
    for section in content['sections']:
        doc.add_heading(section['subtitle'], level=1)
        for paragraph in section['paragraphs']:
            doc.add_paragraph(paragraph, style='BodyText')

    # Salva o documento
    doc.save('output.docx')

# Exemplo de uso
if __name__ == "__main__":
    prompt = input("Por favor, insira o prompt para gerar o conteúdo do documento: ")
    num_sections = int(input("Por favor, insira o número de secções desejado: "))
    max_tokens = int(input("Por favor, insira o número máximo de tokens para o conteúdo gerado: "))
    theme = input("Escolha um tema (corporate/academic): ")

    # Ajustar o prompt para solicitar a divisão em secções
    sectioned_prompt = f"{prompt}\n\nDivida o conteúdo em {num_sections} secções claramente marcadas."

    generated_text = generate_content_from_prompt(sectioned_prompt, max_tokens)

    # Processar o texto gerado em seções
    sections = []
    paragraphs = generated_text.split('\n\n')
    paragraphs_per_section = len(paragraphs) // num_sections

    for i in range(num_sections):
        start = i * paragraphs_per_section
        end = (i + 1) * paragraphs_per_section
        section = {
            'subtitle': f'Section {i + 1}',
            'paragraphs': paragraphs[start:end]
        }
        sections.append(section)

    # Caso haja parágrafos restantes, adicioná-los à última secção
    if len(paragraphs) % num_sections != 0:
        sections[-1]['paragraphs'].extend(paragraphs[num_sections * paragraphs_per_section:])

    content = {
        'title': 'Documento Gerado pelo GPT-4',
        'sections': sections
    }

    create_document(content, theme)
    print("Documento criado com sucesso: output.docx")
